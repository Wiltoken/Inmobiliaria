"""
Convert all Inmobiliaria markdown documentation to Word (.docx) with DIAN-style formatting.

Style reference: DIAN "Lineamientos de Desarrollo SW V1.3"
- Header: blue bar (#1F3864) with document title, version, date
- Footer: page number, classification "Información Pública"
- Version control table
- Professional formatting throughout
"""

import os
import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── Configuration ───────────────────────────────────────────────────────────

DOCS_DIR = Path("/home/userwil/Inmobiliaria/Documentación Inmobiliaria")
OUTPUT_DIR = DOCS_DIR  # Same folder
HU_DIR = DOCS_DIR / "Historias de Usuario Word"

DIAN_BLUE = RGBColor(0x1F, 0x38, 0x64)
DIAN_LIGHT_BLUE = RGBColor(0x2B, 0x57, 0x9A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)
GRAY = RGBColor(0x66, 0x66, 0x66)
DARK_BLUE = RGBColor(0x15, 0x2A, 0x4A)

FONT_NAME = "Calibri"
FONT_SIZE_BODY = Pt(10)
FONT_SIZE_H1 = Pt(18)
FONT_SIZE_H2 = Pt(14)
FONT_SIZE_H3 = Pt(12)
FONT_SIZE_SMALL = Pt(8)
FONT_SIZE_TITLE = Pt(22)

VERSION = "1.0"
VERSION_DATE = datetime.now().strftime("%d/%m/%Y")
PROJECT_NAME = "Inmobiliaria Platform"
CLASIFICACION = "Información Pública"


def create_dian_header(doc, title, version=VERSION, date=VERSION_DATE):
    """Create DIAN-style header with blue bar and document info."""
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False

    # Blue bar at top
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(4)

    # Add a blue rectangle shape (using a table with blue background)
    table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Left cell - Project info
    left = table.cell(0, 0)
    left.width = Inches(4)
    for para in left.paragraphs:
        para.clear()

    p1 = left.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p1.add_run(PROJECT_NAME)
    run.font.size = FONT_SIZE_SMALL
    run.font.color.rgb = GRAY
    run.font.name = FONT_NAME

    p2 = left.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p2.add_run(title)
    run.font.size = Pt(11)
    run.font.color.rgb = DIAN_BLUE
    run.font.name = FONT_NAME
    run.bold = True

    # Right cell - Version + Classification
    right = table.cell(0, 1)
    right.width = Inches(2.5)
    for para in right.paragraphs:
        para.clear()

    p3 = right.paragraphs[0]
    p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p3.add_run(f"Versión {version}")
    run.font.size = FONT_SIZE_SMALL
    run.font.color.rgb = GRAY
    run.font.name = FONT_NAME

    p4 = right.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p4.add_run(date)
    run.font.size = FONT_SIZE_SMALL
    run.font.color.rgb = GRAY
    run.font.name = FONT_NAME

    p5 = right.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p5.add_run(CLASIFICACION)
    run.font.size = FONT_SIZE_SMALL
    run.font.color.rgb = GRAY
    run.font.name = FONT_NAME
    run.italic = True

    # Blue separator line below header
    separator = header.add_paragraph()
    separator.paragraph_format.space_before = Pt(2)
    separator.paragraph_format.space_after = Pt(6)
    run = separator.add_run("─" * 90)
    run.font.size = Pt(4)
    run.font.color.rgb = DIAN_BLUE

    return header


def create_dian_footer(doc):
    """Create DIAN-style footer with page number and classification."""
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False

    # Separator line
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run("─" * 90)
    run.font.size = Pt(4)
    run.font.color.rgb = DIAN_BLUE

    # Footer text
    p2 = footer.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(2)
    run = p2.add_run(f"{PROJECT_NAME} — Documentación Técnica | ")
    run.font.size = FONT_SIZE_SMALL
    run.font.color.rgb = GRAY
    run.font.name = FONT_NAME

    # Page number
    run2 = p2.add_run("Página ")
    run2.font.size = FONT_SIZE_SMALL
    run2.font.color.rgb = GRAY
    run2.font.name = FONT_NAME

    # Add page number field
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run3 = p2.add_run()
    run3._r.append(fldChar1)
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run4 = p2.add_run()
    run4._r.append(instrText)
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run5 = p2.add_run()
    run5._r.append(fldChar2)

    # Classification
    p3 = footer.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p3.add_run(CLASIFICACION)
    run.font.size = FONT_SIZE_SMALL
    run.font.color.rgb = GRAY
    run.font.name = FONT_NAME
    run.italic = True

    return footer


def add_version_table(doc, title, version=VERSION):
    """Add version control table after the title."""
    # Section title
    h = doc.add_heading("Control de Versiones", level=2)
    for run in h.runs:
        run.font.color.rgb = DIAN_BLUE

    table = doc.add_table(rows=2, cols=5, style="Table Grid")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    headers = ["Versión", "Vigencia Desde", "Vigencia Hasta", "Descripción de Cambios", "Autor"]
    for i, header_text in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header_text)
        run.font.size = FONT_SIZE_SMALL
        run.font.bold = True
        run.font.color.rgb = WHITE
        run.font.name = FONT_NAME
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Blue background
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1F3864"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    # Data row
    data = [version, VERSION_DATE, "Vigente", f"Versión inicial — {title}", "Equipo Inmobiliaria"]
    for i, value in enumerate(data):
        cell = table.cell(1, i)
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(value)
        run.font.size = FONT_SIZE_SMALL
        run.font.name = FONT_NAME
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # Spacing
    return table


def add_dian_title_page(doc, title, subtitle=""):
    """Add a DIAN-style cover/title section."""
    # Spacing
    for _ in range(3):
        doc.add_paragraph()

    # Blue bar
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("━" * 60)
    run.font.size = Pt(6)
    run.font.color.rgb = DIAN_BLUE

    # Project name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run(PROJECT_NAME)
    run.font.size = Pt(14)
    run.font.color.rgb = DIAN_BLUE
    run.font.name = FONT_NAME
    run.bold = True

    # Document title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    run = p.add_run(title)
    run.font.size = FONT_SIZE_TITLE
    run.font.color.rgb = DARK_BLUE
    run.font.name = FONT_NAME
    run.bold = True

    if subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(subtitle)
        run.font.size = Pt(12)
        run.font.color.rgb = GRAY
        run.font.name = FONT_NAME

    # Blue bar
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    run = p.add_run("━" * 60)
    run.font.size = Pt(6)
    run.font.color.rgb = DIAN_BLUE

    # Meta info
    meta = [
        ("Versión:", VERSION),
        ("Fecha:", VERSION_DATE),
        ("Clasificación:", CLASIFICACION),
        ("Proyecto:", PROJECT_NAME),
    ]
    table = doc.add_table(rows=len(meta), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(meta):
        cell_l = table.cell(i, 0)
        cell_l.text = ""
        p = cell_l.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(label)
        run.font.size = FONT_SIZE_BODY
        run.font.bold = True
        run.font.name = FONT_NAME

        cell_v = table.cell(i, 1)
        cell_v.text = ""
        p = cell_v.paragraphs[0]
        run = p.add_run(value)
        run.font.size = FONT_SIZE_BODY
        run.font.name = FONT_NAME

    doc.add_page_break()


def style_paragraph(paragraph, font_size=FONT_SIZE_BODY):
    """Apply base styling to a paragraph."""
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.15
    for run in paragraph.runs:
        run.font.name = FONT_NAME
        run.font.size = font_size


def style_heading(heading, level=1):
    """Apply DIAN-style heading formatting."""
    if level == 1:
        heading.runs[0].font.size = FONT_SIZE_H1
    elif level == 2:
        heading.runs[0].font.size = FONT_SIZE_H2
    elif level >= 3:
        heading.runs[0].font.size = FONT_SIZE_H3

    for run in heading.runs:
        run.font.color.rgb = DIAN_BLUE
        run.font.name = FONT_NAME


def md_to_docx(md_path: Path, docx_path: Path, doc_title: str, subtitle=""):
    """Convert a markdown file to a DIAN-styled DOCX."""
    print(f"  📄 {md_path.name} → {docx_path.name}")

    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21.59)  # Letter
    section.page_height = Cm(27.94)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # DIAN style
    create_dian_header(doc, doc_title)
    create_dian_footer(doc)

    # Title page
    add_dian_title_page(doc, doc_title, subtitle)

    # Version table
    add_version_table(doc, doc_title)

    # Read markdown content
    content = md_path.read_text(encoding="utf-8")

    # Simple markdown-to-docx conversion
    lines = content.split("\n")
    in_code_block = False
    in_table = False
    table_rows = []
    skip_frontmatter = True

    for line in lines:
        # Skip YAML frontmatter
        if skip_frontmatter:
            if line.strip() == "---" and len(lines) > 1:
                continue
            skip_frontmatter = False

        stripped = line.strip()

        # Code blocks
        if stripped.startswith("```"):
            in_code_block = not in_code_block
            continue

        if in_code_block:
            p = doc.add_paragraph()
            run = p.add_run(stripped)
            run.font.name = "Consolas"
            run.font.size = Pt(8)
            p.paragraph_format.left_indent = Cm(1)
            continue

        # Headings
        if stripped.startswith("# "):
            h = doc.add_heading(stripped[2:], level=1)
            style_heading(h, 1)
        elif stripped.startswith("## "):
            h = doc.add_heading(stripped[3:], level=2)
            style_heading(h, 2)
        elif stripped.startswith("### "):
            h = doc.add_heading(stripped[4:], level=3)
            style_heading(h, 3)
        elif stripped.startswith("#### "):
            h = doc.add_heading(stripped[5:], level=4)
            style_heading(h, 4)
        # Table rows
        elif stripped.startswith("|") and stripped.endswith("|"):
            if "---" in stripped:
                continue  # separator row
            cells = [c.strip() for c in stripped.split("|")[1:-1]]
            table_rows.append(cells)
        # Horizontal rule
        elif stripped in ("---", "***", "___"):
            p = doc.add_paragraph()
            run = p.add_run("─" * 60)
            run.font.size = Pt(4)
            run.font.color.rgb = DIAN_BLUE
        # Empty line
        elif not stripped:
            if table_rows and not in_table:
                # Flush previous table
                create_table_from_rows(doc, table_rows)
                table_rows = []
            in_table = False
        # Regular text
        else:
            if table_rows:
                create_table_from_rows(doc, table_rows)
                table_rows = []

            # Bold/italic markers
            text = stripped
            p = doc.add_paragraph()
            style_paragraph(p)

            # Simple inline formatting
            parts = re.split(r"(\*\*.*?\*\*|\*.*?\*|`.*?`)", text)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
                    run = p.add_run(part[1:-1])
                    run.italic = True
                elif part.startswith("`") and part.endswith("`"):
                    run = p.add_run(part[1:-1])
                    run.font.name = "Consolas"
                    run.font.size = Pt(8)
                else:
                    run = p.add_run(part)
                run.font.name = FONT_NAME
                run.font.size = FONT_SIZE_BODY

    # Flush remaining table
    if table_rows:
        create_table_from_rows(doc, table_rows)

    doc.save(str(docx_path))


def create_table_from_rows(doc, rows):
    """Create a formatted Word table from parsed markdown rows."""
    if not rows:
        return

    ncols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=ncols, style="Table Grid")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            if j >= ncols:
                break
            cell = table.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            run.font.name = FONT_NAME
            run.font.size = FONT_SIZE_SMALL if i == 0 else FONT_SIZE_BODY
            if i == 0:  # Header row
                run.bold = True
                run.font.color.rgb = WHITE
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1F3864"/>')
                cell._tc.get_or_add_tcPr().append(shading)

    doc.add_paragraph()  # Spacing after table


def create_user_story_template():
    """Create a Word template for user stories matching DIAN format FT-IIT-2006."""
    HU_DIR.mkdir(parents=True, exist_ok=True)
    docx_path = HU_DIR / "Plantilla_Historia_Usuario.docx"

    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21.59)
    section.page_height = Cm(27.94)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    create_dian_header(doc, "Historia de Usuario")
    create_dian_footer(doc)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("HISTORIA DE USUARIO")
    run.font.size = FONT_SIZE_TITLE
    run.font.color.rgb = DARK_BLUE
    run.font.name = FONT_NAME
    run.bold = True

    # Blue bar
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("━" * 50)
    run.font.size = Pt(6)
    run.font.color.rgb = DIAN_BLUE

    # HU Header fields
    fields = [
        ("ID Historia de Usuario:", "HU-INMO-XXX"),
        ("Título:", "[Título descriptivo de la historia]"),
        ("Proyecto:", PROJECT_NAME),
        ("Módulo:", "[Módulo del sistema]"),
        ("Prioridad:", "☐ CRÍTICA  ☐ ALTA  ☐ MEDIA  ☐ BAJA"),
        ("Estimación (horas):", "[N]"),
        ("Sprint:", "[Número]"),
        ("Product Owner:", "[Nombre]"),
        ("Scrum Master:", "[Nombre]"),
    ]

    for label, value in fields:
        p = doc.add_paragraph()
        run = p.add_run(label + " ")
        run.font.bold = True
        run.font.name = FONT_NAME
        run.font.size = FONT_SIZE_BODY
        run = p.add_run(value)
        run.font.name = FONT_NAME
        run.font.size = FONT_SIZE_BODY
        p.paragraph_format.space_after = Pt(4)

    doc.add_paragraph()

    # Sections
    sections = [
        ("1. DESCRIPCIÓN", "Como [rol del usuario]\nQuiero [funcionalidad o acción]\nPara [beneficio o valor de negocio]"),
        ("2. CRITERIOS DE ACEPTACIÓN", "Dado [contexto inicial]\nCuando [acción del usuario]\nEntonces [resultado esperado]\n\nDado [contexto alternativo]\nCuando [acción alternativa]\nEntonces [resultado alternativo]"),
        ("3. REQUISITOS NO FUNCIONALES", "• Seguridad: [requisito]\n• Rendimiento: [requisito]\n• Accesibilidad: [requisito]\n• Disponibilidad: [requisito]"),
        ("4. PROTOTIPOS / MOCKUPS", "[Referencia a mockups en Figma o imágenes adjuntas]\n\n[Insertar imagen aquí si aplica]"),
        ("5. DEPENDENCIAS", "• [Historia o componente del que depende]\n• [API o servicio externo requerido]"),
        ("6. NOTAS / OBSERVACIONES", "[Información adicional relevante para el equipo de desarrollo]"),
        ("7. DEFINITION OF DONE", "☐ Código implementado y revisado\n☐ Pruebas unitarias pasan (>80% coverage)\n☐ Pruebas de integración pasan\n☐ Análisis estático limpio (ruff, mypy)\n☐ Documentación API actualizada (Swagger)\n☐ Desplegable en entorno de pruebas"),
    ]

    for section_title, placeholder in sections:
        h = doc.add_heading(section_title, level=2)
        style_heading(h, 2)
        p = doc.add_paragraph(placeholder)
        p.paragraph_format.left_indent = Cm(1)
        style_paragraph(p)

    # Version table
    doc.add_paragraph()
    doc.add_heading("Control de Versiones", level=2)
    table = doc.add_table(rows=2, cols=4, style="Table Grid")
    headers = ["Versión", "Fecha", "Autor", "Cambios"]
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.bold = True
        run.font.size = FONT_SIZE_SMALL
        run.font.color.rgb = WHITE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1F3864"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    for i in range(4):
        cell = table.cell(1, i)
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run("[—]" if i > 0 else "1.0")
        run.font.size = FONT_SIZE_SMALL
        run.font.name = FONT_NAME
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(str(docx_path))
    print(f"  📄 Plantilla Historia de Usuario → {docx_path.name}")


def convert_all_docs():
    """Convert all markdown documents to DIAN-styled DOCX."""
    print("\n📝 Convirtiendo documentación a Word (formato DIAN)...\n")

    docs_to_convert = [
        ("GOBIERNO_TI.md", "Gobierno TI — Marco de Gobernanza", "ISO 25010 · ISO 27001 · MinTIC"),
        ("SCRUM_METODOLOGIA.md", "Metodología Scrum — Marco Ágil de Desarrollo", "Sprints · Historias de Usuario · Definition of Done"),
        ("CODIGO_LIMPIO_GUIA.md", "Guía de Código Limpio y Buenas Prácticas", "SOLID · Clean Architecture · Refactoring"),
        ("CALIDAD_PRUEBAS_PLAN.md", "Plan de Calidad y Pruebas", "Unitarias · Integración · Rendimiento · Seguridad"),
        ("CUMPLIMIENTO_NORMATIVO.md", "Matriz de Cumplimiento Normativo", "Ley 1581 · Decreto 767 · Fondo Emprender"),
        ("ARCHITECTURE.md", "Arquitectura del Sistema", "Hexagonal · ADRs · Modelo de Dominio"),
        ("INFRASTRUCTURE.md", "Infraestructura Tecnológica", "Servicios · Escalabilidad · Topología"),
        ("API.md", "API Reference — Referencia de Endpoints", "REST · JWT · Swagger"),
        ("DEPLOYMENT.md", "Guía de Despliegue", "Docker · SSL · CI/CD"),
        ("DEVELOPMENT.md", "Guía de Desarrollo", "Setup Local · Tests · Convenciones"),
        ("OPERATIONS.md", "Manual de Operaciones", "Runbooks · Backups · Monitoreo"),
        ("SECURITY.md", "Seguridad de la Información", "JWT · RBAC · OWASP · Ley 1581"),
    ]

    for filename, title, subtitle in docs_to_convert:
        md_path = DOCS_DIR / filename
        if not md_path.exists():
            print(f"  ⚠️ {filename} no encontrado, saltando...")
            continue
        docx_name = filename.replace(".md", ".docx")
        docx_path = DOCS_DIR / docx_name
        md_to_docx(md_path, docx_path, title, subtitle)

    # Create HU template
    create_user_story_template()

    # Create 6 AUTH user stories from existing data
    create_auth_user_stories()

    print(f"\n✅ Conversión completa. {len(docs_to_convert)} documentos + plantilla HU generados.")
    print(f"   Ubicación: {DOCS_DIR}")


def create_auth_user_stories():
    """Create Word docs for the 6 existing AUTH user stories."""
    stories = [
        {
            "id": "AUTH_HU001_REQ001",
            "title": "Sistema de Autenticación y Control de Acceso a Plataforma",
            "role": "Administrador del Sistema",
            "want": "gestionar la autenticación centralizada y el control de acceso basado en roles (RBAC)",
            "benefit": "garantizar que solo usuarios autorizados accedan a los recursos del sistema según su perfil",
            "priority": "CRÍTICA",
            "hours": "80",
            "module": "Autenticación General (AUTH)",
            "acceptance": [
                "Dado un usuario registrado con credenciales válidas\nCuando inicia sesión con usuario y contraseña\nEntonces el sistema valida las credenciales y genera tokens JWT (access + refresh)",
                "Dado un usuario con rol 'admin'\nCuando accede al panel de administración\nEntonces el sistema permite el acceso a funciones administrativas",
                "Dado un usuario con rol 'viewer'\nCuando intenta acceder a funciones de administración\nEntonces el sistema rechaza el acceso con error 403",
            ],
            "nfr": [
                "Seguridad: JWT con HS256, refresh token rotation, blacklist en Redis",
                "Rendimiento: Validación de token < 50ms",
                "Disponibilidad: 99.9% uptime para endpoints de auth",
                "Auditoría: Todo evento de auth registrado en AuditLog",
            ],
        },
        {
            "id": "AUTH_HU002_REQ001",
            "title": "Formulario de Inicio de Sesión y Validación de Credenciales",
            "role": "Usuario registrado",
            "want": "iniciar sesión mediante un formulario seguro con validación de credenciales",
            "benefit": "acceder a las funcionalidades del sistema de forma segura y controlada",
            "priority": "CRÍTICA",
            "hours": "28",
            "module": "Autenticación General (AUTH)",
            "acceptance": [
                "Dado un usuario en la página de login\nCuando ingresa usuario y contraseña correctos\nEntonces el sistema responde con tokens JWT y redirige al dashboard",
                "Dado un usuario en la página de login\nCuando ingresa credenciales incorrectas\nEntonces el sistema muestra mensaje genérico 'Credenciales inválidas' y contador de intentos restantes",
                "Dado un usuario que falla 3 intentos de login\nCuando intenta un cuarto login\nEntonces la cuenta se bloquea y el sistema responde con error 423",
                "Dado el campo de contraseña\nCuando el usuario intenta pegar texto\nEntonces el sistema bloquea la acción de pegado (copy/paste restringido)",
            ],
            "nfr": [
                "Seguridad: Restricción copy/paste en campo contraseña",
                "Usabilidad: Mensajes de error genéricos (sin revelar si el usuario existe)",
                "Rendimiento: Respuesta de validación < 200ms",
            ],
        },
        {
            "id": "AUTH_HU003_REQ001",
            "title": "Integración reCAPTCHA v3 y Protección Anti-Bots",
            "role": "Administrador de Seguridad",
            "want": "integrar reCAPTCHA v3 en el formulario de login",
            "benefit": "proteger la plataforma contra ataques automatizados y bots maliciosos",
            "priority": "CRÍTICA",
            "hours": "12",
            "module": "Autenticación General (AUTH)",
            "acceptance": [
                "Dado un usuario humano que completa el login\nCuando reCAPTCHA v3 evalúa la interacción\nEntonces el score es >= 0.5 y el login procede normalmente",
                "Dado un bot que intenta hacer login automatizado\nCuando reCAPTCHA v3 evalúa la solicitud\nEntonces el score es < 0.5 y el sistema rechaza con error 403",
                "Dado el entorno de desarrollo\nCuando RECAPTCHA_ENABLED=false\nEntonces se usa un mock verifier que aprueba todas las solicitudes",
            ],
            "nfr": [
                "Seguridad: Score mínimo configurable (recaptcha_score_threshold)",
                "Rendimiento: Verificación reCAPTCHA < 500ms",
                "Configurabilidad: Habilitar/deshabilitar por deploy",
            ],
        },
        {
            "id": "AUTH_HU004_REQ001",
            "title": "Recuperación Segura de Contraseña y Desbloqueo de Cuenta",
            "role": "Usuario registrado",
            "want": "recuperar mi contraseña de forma segura mediante un token temporal",
            "benefit": "restaurar el acceso a mi cuenta sin comprometer la seguridad",
            "priority": "ALTA",
            "hours": "16",
            "module": "Autenticación General (AUTH)",
            "acceptance": [
                "Dado un usuario que olvidó su contraseña\nCuando solicita recuperación con su email registrado\nEntonces el sistema genera un token de un solo uso con expiración de 15 minutos",
                "Dado un usuario con token de recuperación válido\nCuando ingresa el token y una nueva contraseña\nEntonces el sistema actualiza la contraseña, marca el token como usado e invalida todas las sesiones activas",
                "Dado un token de recuperación expirado o ya usado\nCuando se intenta usar\nEntonces el sistema rechaza con error 400",
            ],
            "nfr": [
                "Seguridad: Token hash almacenado (no texto plano), single-use, TTL 15 min",
                "Auditoría: Evento de recuperación registrado en AuditLog",
                "Disponibilidad: El token se entrega en la respuesta (token-return-only, sin dependencia SMTP)",
            ],
        },
        {
            "id": "AUTH_HU005_REQ001",
            "title": "Gestión de Sesiones Seguras y Control de Inactividad",
            "role": "Usuario autenticado",
            "want": "que mi sesión se mantenga segura y se cierre automáticamente por inactividad",
            "benefit": "proteger mi cuenta contra accesos no autorizados en equipos compartidos",
            "priority": "ALTA",
            "hours": "16",
            "module": "Autenticación General (AUTH)",
            "acceptance": [
                "Dado un usuario con sesión activa\nCuando realiza acciones en la plataforma\nEntonces el heartbeat de sesión se actualiza en Redis (last_active:{user_id})",
                "Dado un usuario inactivo por 30 minutos\nCuando intenta realizar una acción\nEntonces el sistema invalida la sesión y responde con error 401 (AUTH_SESSION_EXPIRED)",
                "Dado un usuario que hace logout\nCuando confirma la salida\nEntonces el JTI actual se agrega a la blacklist de Redis y el refresh token se revoca",
            ],
            "nfr": [
                "Seguridad: Access token 15 min, Refresh token 7 días, rotación en cada uso",
                "Rendimiento: Verificación de sesión < 10ms (Redis)",
                "Configurabilidad: inactivity_timeout_minutes configurable",
            ],
        },
        {
            "id": "AUTH_HU006_REQ001",
            "title": "Cumplimiento Normativo Colombiano para Acceso Digital",
            "role": "Oficial de Cumplimiento",
            "want": "que la plataforma cumpla con la Ley 1581 de 2012 y normativas colombianas de protección de datos",
            "benefit": "operar legalmente en Colombia y proteger los datos personales de los usuarios",
            "priority": "CRÍTICA",
            "hours": "20",
            "module": "Autenticación General (AUTH)",
            "acceptance": [
                "Dado un nuevo usuario que se registra\nCuando completa el formulario de registro\nEntonces el sistema registra consent_given_at con la fecha y hora del consentimiento",
                "Dado un usuario que solicita sus datos personales\nCuando accede al endpoint de consulta\nEntonces el sistema retorna todos los datos almacenados del titular",
                "Dado un usuario que solicita eliminación de datos\nCuando se procesa la solicitud\nEntonces el sistema elimina o anonimiza los datos personales según la política de retención",
                "Dado un auditor externo\nCuando solicita el reporte de cumplimiento\nEntonces GET /admin/compliance-report retorna estadísticas de consentimiento, datos activos y retención",
            ],
            "nfr": [
                "Legal: Cumplimiento Ley 1581/2012, Decreto 1377/2013",
                "Auditoría: Retención de audit_logs configurable (default 365 días)",
                "Seguridad: Datos personales cifrados en tránsito (TLS 1.3) y en reposo",
                "Reportes: Endpoint de compliance report para evidencia de cumplimiento",
            ],
        },
    ]

    for story in stories:
        doc = Document()

        section = doc.sections[0]
        section.page_width = Cm(21.59)
        section.page_height = Cm(27.94)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

        create_dian_header(doc, f"Historia de Usuario — {story['id']}")
        create_dian_footer(doc)

        # Title
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("HISTORIA DE USUARIO")
        run.font.size = FONT_SIZE_TITLE
        run.font.color.rgb = DARK_BLUE
        run.font.name = FONT_NAME
        run.bold = True

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("━" * 50)
        run.font.size = Pt(6)
        run.font.color.rgb = DIAN_BLUE

        # Fields
        fields = [
            ("ID:", story["id"]),
            ("Título:", story["title"]),
            ("Proyecto:", PROJECT_NAME),
            ("Módulo:", story["module"]),
            ("Prioridad:", story["priority"]),
            ("Estimación (horas):", story["hours"]),
        ]
        for label, value in fields:
            p = doc.add_paragraph()
            run = p.add_run(label + " ")
            run.font.bold = True
            run.font.name = FONT_NAME
            run.font.size = FONT_SIZE_BODY
            run = p.add_run(value)
            run.font.name = FONT_NAME
            run.font.size = FONT_SIZE_BODY

        doc.add_paragraph()

        # Description
        h = doc.add_heading("1. DESCRIPCIÓN", level=2)
        style_heading(h, 2)
        desc = f"Como {story['role']}\nQuiero {story['want']}\nPara {story['benefit']}"
        p = doc.add_paragraph(desc)
        style_paragraph(p)
        p.paragraph_format.left_indent = Cm(1)

        # Acceptance criteria
        h = doc.add_heading("2. CRITERIOS DE ACEPTACIÓN", level=2)
        style_heading(h, 2)
        for i, ac in enumerate(story["acceptance"], 1):
            p = doc.add_paragraph(f"Criterio {i}:\n{ac}")
            style_paragraph(p)
            p.paragraph_format.left_indent = Cm(1)

        # NFR
        h = doc.add_heading("3. REQUISITOS NO FUNCIONALES", level=2)
        style_heading(h, 2)
        for nfr in story["nfr"]:
            p = doc.add_paragraph(f"• {nfr}")
            style_paragraph(p)
            p.paragraph_format.left_indent = Cm(1)

        # DoD
        h = doc.add_heading("4. DEFINITION OF DONE", level=2)
        style_heading(h, 2)
        dod_items = [
            "Código implementado y revisado (PR aprobado)",
            "Pruebas unitarias e integración pasan",
            "Análisis estático limpio (ruff, mypy, bandit)",
            "Documentación API actualizada (Swagger)",
            "Cumple criterios de aceptación descritos",
        ]
        for item in dod_items:
            p = doc.add_paragraph(f"☐ {item}")
            style_paragraph(p)
            p.paragraph_format.left_indent = Cm(1)

        # Version table
        doc.add_paragraph()
        table = doc.add_table(rows=2, cols=4, style="Table Grid")
        headers = ["Versión", "Fecha", "Autor", "Cambios"]
        for i, h_text in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(h_text)
            run.font.bold = True
            run.font.size = FONT_SIZE_SMALL
            run.font.color.rgb = WHITE
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1F3864"/>')
            cell._tc.get_or_add_tcPr().append(shading)

        vdata = ["1.0", VERSION_DATE, "Equipo Inmobiliaria", "Versión inicial"]
        for i, val in enumerate(vdata):
            cell = table.cell(1, i)
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = FONT_SIZE_SMALL
            run.font.name = FONT_NAME
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        docx_path = HU_DIR / f"{story['id']} — {story['title'][:80]}.docx"
        doc.save(str(docx_path))
        print(f"  📄 {story['id']} → {docx_path.name}")

    print(f"\n  ✅ {len(stories)} historias de usuario generadas en {HU_DIR}")


if __name__ == "__main__":
    convert_all_docs()
